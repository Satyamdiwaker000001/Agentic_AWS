#!/usr/bin/env python3
"""Create a sample DOCX file for testing."""

from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('Machine Learning Overview', 0)

# Add paragraphs
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'Machine learning is a branch of artificial intelligence that enables systems '
    'to learn and improve from experience without being explicitly programmed.'
)

doc.add_heading('Types of Machine Learning', level=1)
doc.add_paragraph(
    'There are three main types of machine learning: supervised learning, '
    'unsupervised learning, and reinforcement learning.'
)

doc.add_heading('Applications', level=1)
applications = [
    'Healthcare: Diagnosis and prediction',
    'Finance: Fraud detection',
    'Retail: Recommendations',
    'Transportation: Autonomous systems'
]
for app in applications:
    doc.add_paragraph(app, style='List Bullet')

# Add a table
doc.add_heading('Comparison', level=1)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Type'
header_cells[1].text = 'Description'

data = [
    ('Supervised', 'Uses labeled data'),
    ('Unsupervised', 'Finds patterns in data'),
    ('Reinforcement', 'Learning through feedback')
]

for i, (typ, desc) in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = typ
    row_cells[1].text = desc

# Save the document
doc.save('uploads/sample.docx')
print("✓ Created uploads/sample.docx")
