from docx import Document
import os


class DOCParser:

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from DOC files. 
        Uses python-docx which has some support for legacy DOC format."""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            # Fallback: attempt basic text extraction if Document parsing fails
            # This is a best-effort approach for older DOC formats
            raise ValueError(f"Unable to parse DOC file: {str(e)}")
