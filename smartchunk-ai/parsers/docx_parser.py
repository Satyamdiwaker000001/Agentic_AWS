from docx import Document


class DOCXParser:

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from DOCX files."""
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
