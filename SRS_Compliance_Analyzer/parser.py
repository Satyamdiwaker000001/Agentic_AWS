# parser.py
# File parsing module to extract structured text and tables from PDF, DOCX, and TXT files.

import os
import re
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(filepath):
    """
    Extract raw text content from a PDF file using pypdf.
    """
    text = ""
    try:
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(filepath):
    """
    Extract raw text content from a DOCX file using python-docx.
    """
    text = []
    try:
        doc = Document(filepath)
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {e}")
    return "\n".join(text)

def extract_text_from_txt(filepath):
    """
    Extract raw text content from a TXT file, trying common encodings.
    """
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file using standard encodings (UTF-8, Latin-1, CP1252).")

def parse_document(filepath):
    """
    Determine file extension and extract raw text (retains backwards compatibility).
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Please upload a .pdf, .docx, or .txt file.")

def extract_blocks_from_docx(filepath):
    """
    Sequentially extract paragraphs and tables from DOCX, maintaining headings association.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        doc = Document(filepath)
    except Exception as e:
        raise ValueError(f"Error reading DOCX blocks: {e}")

    blocks = []
    current_heading = None
    heading_regex = re.compile(r'^(\d+(\.\d+)*)\s+[A-Za-z]', re.IGNORECASE)

    for child in doc.element.body:
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue

            is_head = False
            if para.style is not None:
                try:
                    is_head = para.style.name.startswith("Heading")
                except AttributeError:
                    pass
            is_head = is_head or bool(heading_regex.match(text))

            if is_head:
                current_heading = text
                blocks.append({
                    "text": text,
                    "type": "heading",
                    "preceding_heading": current_heading,
                    "headers": []
                })
            else:
                blocks.append({
                    "text": text,
                    "type": "paragraph",
                    "preceding_heading": current_heading,
                    "headers": []
                })

        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            table_rows = []
            headers = []

            for r_idx, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                # Filter duplicates from merged cells
                cleaned_cells = []
                for cell in row_cells:
                    if not cleaned_cells or cleaned_cells[-1] != cell:
                        cleaned_cells.append(cell)

                if cleaned_cells:
                    row_str = " | ".join(cleaned_cells)
                    table_rows.append(row_str)
                    if r_idx == 0:
                        headers = [c.lower() for c in cleaned_cells]

            if table_rows:
                blocks.append({
                    "text": "\n".join(table_rows),
                    "type": "table",
                    "preceding_heading": current_heading,
                    "headers": headers
                })

    return blocks

def extract_blocks_from_plaintext(text):
    """
    Extract structured blocks from flat text (PDF and TXT), parsing headings dynamically.
    """
    blocks = []
    paras = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    current_heading = None
    heading_regex = re.compile(r'^(\d+(\.\d+)*)\s+[A-Za-z]', re.IGNORECASE)

    for para in paras:
        # Peek at first line of the block
        lines = para.split('\n')
        first_line = lines[0].strip()
        is_head = bool(heading_regex.match(first_line)) and len(first_line) < 100

        if is_head:
            current_heading = first_line
            blocks.append({
                "text": first_line,
                "type": "heading",
                "preceding_heading": current_heading,
                "headers": []
            })
            remaining = "\n".join(lines[1:]).strip()
            if remaining:
                blocks.append({
                    "text": remaining,
                    "type": "paragraph",
                    "preceding_heading": current_heading,
                    "headers": []
                })
        else:
            blocks.append({
                "text": para,
                "type": "paragraph",
                "preceding_heading": current_heading,
                "headers": []
            })

    return blocks

def parse_document_to_blocks(filepath):
    """
    Parser entrypoint returning sequential block dictionaries with layout info.
    """
    import traceback
    try:
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".docx":
            return extract_blocks_from_docx(filepath)
        else:
            # Fall back to raw text extraction then chunk blocks
            raw_text = parse_document(filepath)
            return extract_blocks_from_plaintext(raw_text)
    except Exception as e:
        tb_str = traceback.format_exc()
        print(f"[ERROR] Exception in parse_document_to_blocks:\n{tb_str}")
        raise ValueError(f"Parser error: {str(e)}\n\nTraceback:\n{tb_str}")

